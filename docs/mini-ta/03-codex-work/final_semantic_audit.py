from collections import Counter
from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn


root = Path(__file__).resolve().parents[3]
path = root / "docs" / "mini-ta" / "04-output" / "AKSARA_LAPORAN_FINAL.docx"
document = Document(str(path))

expected_title = (
    "AKSARA (Authenticated Key-based Secure Autonomous Relay Architecture): "
    "Chat Terminal Tanpa Server — Implementasi dan Evaluasi Keamanan Protokol "
    "Noise_IK, Siklus Hidup Kunci, dan Threat Model"
)
members = {
    "Andika Aryansyach Fauzan": "2322101878",
    "Mahendra Nur Hidayat": "2322101937",
    "Rafi Putra Fadlurrahman": "2322101963",
}

paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
table_texts = [
    cell.text
    for table in document.tables
    for row in table.rows
    for cell in row.cells
]
full_text = "\n".join(paragraph_texts + table_texts)

assert expected_title in full_text, "Judul final tidak cocok"
for name, student_id in members.items():
    assert name in full_text and student_id in full_text, f"Identitas anggota hilang: {name}"

forbidden = {
    "CARAKA",
    "CLAMP",
    "Compact Lightweight Authenticated Mesh Protocol",
    "Ascon-MAC",
    "Ascon-Hash256",
    "Epidemic Sync",
    "Controlled Flooding",
    "[PERLU KONFIRMASI]",
    "NEEDS_CONFIRMATION",
    "[[CITE:",
    "Daftar isi diperbarui saat dokumen dibuka.",
    "Daftar gambar diperbarui saat dokumen dibuka.",
    "Daftar tabel diperbarui saat dokumen dibuka.",
}
found_forbidden = sorted(term for term in forbidden if term.casefold() in full_text.casefold())
assert not found_forbidden, f"Konten terlarang/placeholder ditemukan: {found_forbidden}"

heading_counts = Counter(paragraph.style.name for paragraph in document.paragraphs)
assert heading_counts["Heading 1"] == 7, heading_counts
assert heading_counts["Heading 2"] == 34, heading_counts

figure_captions = [
    paragraph.text
    for paragraph in document.paragraphs
    if paragraph.style.name == "Caption Figure"
]
table_captions = [
    paragraph.text
    for paragraph in document.paragraphs
    if paragraph.style.name == "Caption Table"
]
assert len(figure_captions) == 18, len(figure_captions)
assert len(table_captions) == 13, len(table_captions)
assert len(document.inline_shapes) == 19, len(document.inline_shapes)
assert len(document.tables) == 13, len(document.tables)

figure_by_chapter = Counter(int(re.match(r"Gambar (\d+)\.", caption).group(1)) for caption in figure_captions)
table_by_chapter = Counter(int(re.match(r"Tabel (\d+)\.", caption).group(1)) for caption in table_captions)
assert figure_by_chapter == Counter({4: 17, 5: 1}), figure_by_chapter
assert table_by_chapter == Counter({1: 1, 2: 4, 3: 1, 4: 5, 5: 2}), table_by_chapter

bibliography_index = paragraph_texts.index("DAFTAR PUSTAKA")
body_text = "\n".join(paragraph_texts[:bibliography_index])
bibliography_texts = paragraph_texts[bibliography_index + 1 :]
bibliography_numbers = [
    int(match.group(1))
    for text in bibliography_texts
    if (match := re.match(r"^\[(\d+)\]", text))
]
assert bibliography_numbers == list(range(1, 41)), bibliography_numbers
cited_numbers = {int(value) for value in re.findall(r"\[(\d+)\]", body_text)}
assert cited_numbers == set(range(1, 41)), sorted(cited_numbers)

page_fields = 0
for section in document.sections:
    for footer in (section.footer, section.first_page_footer, section.even_page_footer):
        for field in footer._element.findall(".//" + qn("w:instrText")):
            if (field.text or "").strip() == "PAGE":
                page_fields += 1
assert page_fields >= 2, page_fields

footer_text = " ".join(
    cell.text
    for section in document.sections
    for footer in (section.footer, section.first_page_footer, section.even_page_footer)
    for table in footer.tables
    for row in table.rows
    for cell in row.cells
)
assert "Politeknik Siber dan Sandi Negara" in footer_text

settings = document.settings.element
update_fields = settings.find(qn("w:updateFields"))
assert update_fields is not None
assert update_fields.get(qn("w:val")) == "false"

restart_sections = []
for index, section in enumerate(document.sections, start=1):
    page_number_type = section._sectPr.find(qn("w:pgNumType"))
    if page_number_type is not None and page_number_type.get(qn("w:start")) is not None:
        restart_sections.append(index)
assert restart_sections == [2, 3], restart_sections

print(f"Document={path}")
print(f"Headings=Heading1:{heading_counts['Heading 1']} Heading2:{heading_counts['Heading 2']}")
print(f"Figures={len(figure_captions)} + cover-logo=1; by-chapter={dict(sorted(figure_by_chapter.items()))}")
print(f"Tables={len(table_captions)}; by-chapter={dict(sorted(table_by_chapter.items()))}")
print("References=40/40 cited and listed")
print("Confirmations=0")
print(f"Page-number restarts={restart_sections}")
print("Semantic audit=PASS")
