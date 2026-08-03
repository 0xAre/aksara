from __future__ import annotations

import re
import shutil
from collections import OrderedDict
from pathlib import Path
from zipfile import ZipFile

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
MINI = ROOT / "docs" / "mini-ta"
PREP = MINI / "01-claude-preparation"
WORK = MINI / "03-codex-work"
OUTPUT = MINI / "04-output"
REFERENCE = MINI / "00-template" / "Cetak TA_rev3.docx"
WORKING_TEMPLATE = WORK / "template-distill" / "Cetak_TA_rev3_working.docx"
LOGO = WORK / "template-distill" / "logo-poltekssn.png"
FINAL = OUTPUT / "AKSARA_LAPORAN_FINAL.docx"

TITLE = (
    "AKSARA (Authenticated Key-based Secure Autonomous Relay Architecture): "
    "Chat Terminal Tanpa Server — Implementasi dan Evaluasi Keamanan Protokol "
    "Noise_IK, Siklus Hidup Kunci, dan Threat Model"
)

MEMBERS = [
    ("Andika Aryansyach Fauzan", "2322101878"),
    ("Mahendra Nur Hidayat", "2322101937"),
    ("Rafi Putra Fadlurrahman", "2322101963"),
]

CHAPTERS = {
    1: ("I", "PENDAHULUAN"),
    2: ("II", "KAJIAN PUSTAKA"),
    3: ("III", "METODOLOGI PENELITIAN"),
    4: ("IV", "PERANCANGAN DAN IMPLEMENTASI"),
    5: ("V", "PENGUJIAN DAN ANALISIS"),
    6: ("VI", "PENUTUP"),
}


def local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def strip_bib_value(value: str) -> str:
    value = value.strip().rstrip(",").strip()
    while len(value) >= 2 and (
        (value[0] == "{" and value[-1] == "}")
        or (value[0] == '"' and value[-1] == '"')
    ):
        value = value[1:-1].strip()
    value = value.replace(r"\&", "&")
    value = value.replace(r"\texttt", "")
    value = value.replace("{", "").replace("}", "")
    value = value.replace("--", "–")
    return value


def parse_bib(path: Path) -> OrderedDict[str, dict[str, str]]:
    text = path.read_text(encoding="utf-8")
    entries: OrderedDict[str, dict[str, str]] = OrderedDict()
    i = 0
    while True:
        match = re.search(r"@(\w+)\s*\{\s*([^,\s]+)\s*,", text[i:])
        if not match:
            break
        entry_type = match.group(1).lower()
        key = match.group(2)
        start = i + match.end()
        depth = 1
        j = start
        while j < len(text) and depth:
            if text[j] == "{":
                depth += 1
            elif text[j] == "}":
                depth -= 1
            j += 1
        body = text[start : j - 1]
        fields: dict[str, str] = {"ENTRYTYPE": entry_type, "ID": key}
        pos = 0
        while pos < len(body):
            field_match = re.search(r"(?m)^\s*(\w+)\s*=\s*", body[pos:])
            if not field_match:
                break
            name = field_match.group(1).lower()
            value_start = pos + field_match.end()
            if value_start >= len(body):
                break
            opener = body[value_start]
            if opener == "{":
                d = 1
                k = value_start + 1
                while k < len(body) and d:
                    d += body[k] == "{"
                    d -= body[k] == "}"
                    k += 1
                raw = body[value_start:k]
            elif opener == '"':
                k = value_start + 1
                while k < len(body):
                    if body[k] == '"' and body[k - 1] != "\\":
                        k += 1
                        break
                    k += 1
                raw = body[value_start:k]
            else:
                k = body.find(",", value_start)
                if k < 0:
                    k = len(body)
                raw = body[value_start:k]
            fields[name] = strip_bib_value(raw)
            pos = k
        entries[key] = fields
        i = j
    return entries


class CitationManager:
    def __init__(self, entries: OrderedDict[str, dict[str, str]]) -> None:
        self.entries = entries
        self.order: list[str] = []
        self.number: dict[str, int] = {}

    def cite(self, keys: list[str]) -> str:
        nums: list[int] = []
        for key in keys:
            if key not in self.entries:
                raise KeyError(f"Citekey tidak ada: {key}")
            if key not in self.number:
                self.order.append(key)
                self.number[key] = len(self.order)
            nums.append(self.number[key])
        return " " + ", ".join(f"[{n}]" for n in nums)

    def uncited(self) -> list[str]:
        return [key for key in self.entries if key not in self.number]


BIB = parse_bib(PREP / "references" / "REFERENCES.bib")
CITES = CitationManager(BIB)


def set_run_font(run, name: str = "Times New Roman", size: float = 12, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, total_cm: float, weights: list[float]):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_twips = int(total_cm / 2.54 * 1440)
    s = sum(weights)
    widths = [int(total_twips * w / s) for w in weights]
    widths[-1] += total_twips - sum(widths)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(widths[idx] / 1440 * 2.54)


def add_field(paragraph, instruction: str, cached_text: str = ""):
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)
    if cached_text:
        cached = paragraph.add_run(cached_text)
        set_run_font(cached)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def set_page_numbering(section, fmt: str, start: int | None = None):
    sect_pr = section._sectPr
    node = sect_pr.find(qn("w:pgNumType"))
    if node is None:
        node = OxmlElement("w:pgNumType")
        sect_pr.append(node)
    node.set(qn("w:fmt"), fmt)
    if start is not None:
        node.set(qn("w:start"), str(start))
    elif qn("w:start") in node.attrib:
        del node.attrib[qn("w:start")]


def clear_part(part):
    element = part._element
    for child in list(element):
        element.remove(child)
    element.append(OxmlElement("w:p"))


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def setup_footer(section, roman: bool = False):
    section.footer.is_linked_to_previous = False
    clear_part(section.footer)
    width = Cm(14)
    table = section.footer.add_table(rows=1, cols=3, width=width)
    remove_table_borders(table)
    set_table_geometry(table, 14, [1, 1, 1.7])
    row = table.rows[0]
    row.cells[0].text = ""

    center = row.cells[1].paragraphs[0]
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(center, "PAGE", "i" if roman else "1")
    for run in center.runs:
        set_run_font(run, "Times New Roman", 12)

    right = row.cells[2].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right.add_run("Politeknik Siber dan Sandi Negara")
    set_run_font(run, "Arial", 10, bold=True)
    for cell in row.cells:
        set_cell_margins(cell, 0, 0, 0, 0)


def set_section_geometry(section, landscape: bool = False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.gutter = Cm(0)
    section.different_first_page_header_footer = False


def set_outline_level(style, level: int):
    p_pr = style._element.get_or_add_pPr()
    node = p_pr.find(qn("w:outlineLvl"))
    if node is None:
        node = OxmlElement("w:outlineLvl")
        p_pr.append(node)
    node.set(qn("w:val"), str(level))


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(13.8)
    normal.paragraph_format.first_line_indent = Cm(0)
    normal.paragraph_format.left_indent = Cm(0)
    normal.paragraph_format.right_indent = Cm(0)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(36)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True
    set_outline_level(h1, 0)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(12)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True
    set_outline_level(h2, 1)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(12)
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.keep_with_next = True
    set_outline_level(h3, 2)

    def create_or_get(name: str, base: str):
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base]
        return style

    front = create_or_get("Front Matter Heading", "Normal")
    front.font.name = "Times New Roman"
    front._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    front._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    front.font.size = Pt(12)
    front.font.bold = True
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.space_after = Pt(24)
    front.paragraph_format.keep_with_next = True

    abstract = create_or_get("Abstract Body", "Normal")
    abstract.paragraph_format.line_spacing = 1.0
    abstract.paragraph_format.space_after = Pt(12)

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(12)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(13.8)

    cap_fig = create_or_get("Caption Figure", "Caption")
    cap_fig.paragraph_format.keep_with_next = False

    cap_tbl = create_or_get("Caption Table", "Caption")
    cap_tbl.paragraph_format.keep_with_next = True
    cap_tbl.paragraph_format.space_after = Pt(6)

    table_text = create_or_get("Table Text", "Normal")
    table_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table_text.paragraph_format.line_spacing = 1.0
    table_text.paragraph_format.space_after = Pt(0)
    table_text.font.size = Pt(12)

    bibliography = create_or_get("Bibliography", "Normal")
    bibliography.paragraph_format.line_spacing = 1.0
    bibliography.paragraph_format.space_after = Pt(6)
    bibliography.paragraph_format.left_indent = Cm(1)
    bibliography.paragraph_format.first_line_indent = Cm(-1)


def clear_document_body(doc: Document):
    body = doc._element.body
    for child in list(body):
        if local_name(child.tag) != "sectPr":
            body.remove(child)


TOKEN_RE = re.compile(r"(\[\[CITE:[^\]]+\]\]|\*[^*]+\*|`[^`]+`)")


def add_markup_runs(paragraph, text: str):
    for token in TOKEN_RE.split(text):
        if not token:
            continue
        if token.startswith("[[CITE:"):
            keys = [x.strip() for x in token[7:-2].split(",") if x.strip()]
            run = paragraph.add_run(CITES.cite(keys))
            set_run_font(run)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, italic=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Courier New", 10.5)
        else:
            run = paragraph.add_run(token)
            set_run_font(run)


def add_para(doc: Document, text: str, style: str = "Normal", align=None, keep=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True
    add_markup_runs(p, text)
    return p


def create_numbering(doc: Document):
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
        if x.get(qn("w:abstractNumId"))
    ]
    existing_num = [
        int(x.get(qn("w:numId")))
        for x in numbering.findall(qn("w:num"))
        if x.get(qn("w:numId"))
    ]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1
    result = {}
    for fmt, text, key in (("decimal", "%1.", "decimal"), ("lowerLetter", "%1.", "letter")):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abs))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl.append(p_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(next_abs))
        num.append(abs_id)
        numbering.append(num)
        result[key] = next_num
        next_abs += 1
        next_num += 1
    return result


def add_list_item(doc: Document, text: str, num_id: int):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    add_markup_runs(p, text)
    return p


def add_lettered_item(doc: Document, text: str, index: int):
    """Add a deterministic a., b., ... item that always restarts per group."""
    value = index + 1
    letters = ""
    while value:
        value, remainder = divmod(value - 1, 26)
        letters = chr(ord("a") + remainder) + letters
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    prefix = p.add_run(f"{letters}.\t")
    set_run_font(prefix)
    add_markup_runs(p, text)
    return p


def add_front_heading(doc: Document, text: str):
    return add_para(doc, text, "Front Matter Heading", WD_ALIGN_PARAGRAPH.CENTER)


def add_chapter(doc: Document, number: int, intro: str):
    roman, title = CHAPTERS[number]
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(f"BAB {roman}\n{title}")
    for run in p.runs:
        set_run_font(run, bold=True)
    add_para(doc, intro)


def add_subheading(doc: Document, number: int, sub: int, title: str):
    roman = CHAPTERS[number][0]
    return add_para(doc, f"{roman}.{sub} {title.upper()}", "Heading 2")


def add_child_heading(doc: Document, number: int, sub: int, child: int, title: str):
    roman = CHAPTERS[number][0]
    return add_para(doc, f"{roman}.{sub}.{child} {title}", "Heading 3")


def parse_first_markdown_table(path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    for idx in range(len(lines) - 2):
        if lines[idx].lstrip().startswith("|") and re.match(r"^\s*\|?\s*:?-+", lines[idx + 1]):
            rows = []
            j = idx
            while j < len(lines) and lines[j].lstrip().startswith("|"):
                if j != idx + 1:
                    row = [x.strip() for x in lines[j].strip().strip("|").split("|")]
                    rows.append(row)
                j += 1
            return rows
    raise ValueError(f"Tabel Markdown tidak ditemukan: {path}")


def clean_table_text(value: str) -> str:
    value = value.replace("**", "").replace("`", "")
    value = value.replace("<br>", "\n").replace("<br/>", "\n")
    for key in sorted(BIB.keys(), key=len, reverse=True):
        if re.search(rf"\b{re.escape(key)}\b", value):
            value = re.sub(rf"\b{re.escape(key)}\b", CITES.cite([key]).strip(), value)
    return value


TABLE_WEIGHTS = {
    2: [1, 3],
    3: [1, 2, 3],
    4: [1.2, 2.3, 3.7, 2.8],
    5: [0.8, 2.2, 3.3, 2.2, 2.0],
    6: [0.6, 2.0, 1.4, 1.1, 1.1, 2.4],
    7: [0.6, 1.5, 2.0, 1.5, 1.4, 2.0, 1.2],
    8: [0.6, 1.2, 1.6, 1.7, 1.3, 1.4, 2.0, 1.1],
}


def add_table(doc: Document, caption: str, rows: list[list[str]], landscape: bool = False):
    if landscape:
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        set_section_geometry(sec, landscape=True)
        sec.footer.is_linked_to_previous = True
    add_para(doc, caption, "Caption Table", WD_ALIGN_PARAGRAPH.CENTER, keep=True)
    cols = len(rows[0])
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    for i, value in enumerate(rows[0]):
        cell = table.rows[0].cells[i]
        cell.text = clean_table_text(value)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "D9E2F3")
        for p in cell.paragraphs:
            p.style = doc.styles["Table Text"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, bold=True)
    set_repeat_table_header(table.rows[0])

    for source_row in rows[1:]:
        if len(source_row) != cols:
            source_row = (source_row + [""] * cols)[:cols]
        cells = table.add_row().cells
        for i, value in enumerate(source_row):
            cells[i].text = clean_table_text(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.style = doc.styles["Table Text"]
                if len(value) <= 18 and i != 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run)
            set_cell_margins(cells[i])
    total = 22.7 if landscape else 14.0
    if caption.startswith("Tabel 4.1 "):
        # Give source paths enough room so the final evidence cell does not
        # spill as a one-line fragment onto a third landscape page.
        weights = [0.7, 2.1, 3.4, 3.2, 1.2]
    else:
        weights = TABLE_WEIGHTS.get(cols, [1] * cols)
    set_table_geometry(table, total, weights)
    if landscape:
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        set_section_geometry(sec, landscape=False)
        sec.footer.is_linked_to_previous = True
    return table


def add_table_from_file(doc: Document, caption: str, filename: str, landscape=True):
    rows = parse_first_markdown_table(PREP / "tables" / filename)
    return add_table(doc, caption, rows, landscape=landscape)


def add_figure(doc: Document, path: Path, caption: str):
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as image:
        width_px, height_px = image.size
    max_w = 14.0
    max_h = 17.0
    ratio = width_px / height_px
    width_cm = min(max_w, max_h * ratio)
    height_cm = width_cm / ratio
    if height_cm > max_h:
        height_cm = max_h
        width_cm = height_cm * ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
    add_para(doc, caption, "Caption Figure", WD_ALIGN_PARAGRAPH.CENTER)


def custom_packet_table():
    return [
        ["Struktur", "Field", "Ukuran", "Deskripsi"],
        ["Invite code", "Ed25519 public key", "32 byte", "Bagian pertama sebelum base64url tanpa padding"],
        ["Invite code", "X25519 public key", "32 byte", "Bagian kedua; gabungan 64 byte menghasilkan 86 karakter"],
        ["Invite code", "Suffix onion", "Variabel", "Opsional @<onion-address>"],
        ["Vault identitas", "Salt", "16 byte", "Salt Argon2id acak per seal()"],
        ["Vault identitas", "Nonce", "12 byte", "Nonce ChaCha20-Poly1305 acak"],
        ["Vault identitas", "Ciphertext", "64 byte", "Ed25519 secret key || X25519 secret key"],
        ["Vault identitas", "Tag", "16 byte", "Tag Poly1305; total vault 108 byte"],
        ["Frame transport", "Length prefix", "2 byte", "Big-endian, panjang payload berikutnya"],
        ["Frame transport", "Payload", "≤ 65.535 byte", "Satu frame memuat satu pesan Noise"],
        ["Payload sesi", "TYPE_TEXT", "1 byte + body", "Tag 0x00 diikuti UTF-8 pesan"],
        ["Payload sesi", "TYPE_BLUR", "2 byte", "Tag 0x01 dan satu byte status"],
        ["Payload sesi", "TYPE_PING", "1 byte", "Tag 0x02 tanpa body"],
        ["Contact store", "Nonce", "12 byte", "Nonce acak per save_contacts()"],
        ["Contact store", "Ciphertext + tag", "Variabel", "Data kontak terenkripsi tanpa magic header"],
    ]


def author_ieee(author: str) -> str:
    if not author:
        return ""
    if author.startswith("{") or "Developers" in author or "contributors" in author or "Project" in author or "Institute" in author:
        return author.strip("{}")
    rendered = []
    for person in author.split(" and "):
        person = person.strip()
        parts = person.split()
        if len(parts) == 1:
            rendered.append(parts[0])
            continue
        initials = " ".join(f"{p[0]}." for p in parts[:-1] if p)
        rendered.append(f"{initials} {parts[-1]}")
    if len(rendered) == 1:
        return rendered[0]
    if len(rendered) == 2:
        return " and ".join(rendered)
    return ", ".join(rendered[:-1]) + ", and " + rendered[-1]


def ieee_entry(number: int, entry: dict[str, str]) -> str:
    authors = author_ieee(entry.get("author", ""))
    title = entry.get("title", "")
    year = entry.get("year", "")
    typ = entry.get("ENTRYTYPE", "")
    pieces = [f"[{number}]"]
    if authors:
        pieces.append(authors + ",")
    pieces.append(f'“{title},”')
    if typ == "article":
        if entry.get("journal"):
            pieces.append(f"*{entry['journal']}*,")
        if entry.get("volume"):
            pieces.append(f"vol. {entry['volume']},")
        if entry.get("number"):
            pieces.append(f"no. {entry['number']},")
        if entry.get("pages"):
            pieces.append(f"pp. {entry['pages']},")
    elif typ == "inproceedings":
        if entry.get("booktitle"):
            pieces.append(f"in *{entry['booktitle']}*,")
        if entry.get("pages"):
            pieces.append(f"pp. {entry['pages']},")
    else:
        if entry.get("howpublished"):
            pieces.append(entry["howpublished"] + ",")
        elif entry.get("institution"):
            pieces.append(entry["institution"] + ",")
    if year:
        pieces.append(year + ".")
    if entry.get("doi"):
        pieces.append(f"doi: {entry['doi']}.")
    if entry.get("url"):
        pieces.append(f"[Online]. Available: {entry['url']}")
    return " ".join(pieces)


def set_update_fields(doc: Document):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_cover(doc: Document):
    section = doc.sections[0]
    set_section_geometry(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_part(section.header)
    clear_part(section.footer)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("TUGAS MATA KULIAH\nIMPLEMENTASI KRIPTOGRAFI")
    set_run_font(run, bold=True)

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        p.add_run().add_picture(str(LOGO), width=Cm(4.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(TITLE)
    set_run_font(run, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Disusun oleh:")
    set_run_font(run)
    for name, nim in MEMBERS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{name} ({nim})")
        set_run_font(run, bold=True)

    for _ in range(2):
        doc.add_paragraph()
    for text in (
        "PROGRAM STUDI REKAYASA SISTEM KRIPTOGRAFI",
        "POLITEKNIK SIBER DAN SANDI NEGARA",
        "2026",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, bold=True)


def add_front_matter(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(section)
    set_page_numbering(section, "lowerRoman", 1)
    setup_footer(section, roman=True)

    add_front_heading(doc, "ABSTRAK")
    p = doc.add_paragraph(style="Abstract Body")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    set_run_font(r, bold=True)
    p = doc.add_paragraph(style="Abstract Body")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_markup_runs(p, "Oleh")
    for name, nim in MEMBERS:
        p = doc.add_paragraph(style="Abstract Body")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_markup_runs(p, f"{name} ({nim})")
    p = doc.add_paragraph(style="Abstract Body")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_markup_runs(p, "Program Studi Rekayasa Sistem Kriptografi")

    abstract = (
        "AKSARA merupakan aplikasi chat terminal peer-to-peer terenkripsi untuk dua pihak yang "
        "beroperasi tanpa server perantara. Penelitian deskriptif-evaluatif ini mendokumentasikan "
        "implementasi protokol keamanan AKSARA versi 0.2.1, menganalisis siklus hidup kunci, "
        "menyusun threat model, serta mengevaluasi perilaku correctness, rejection, dan performa "
        "secara terbatas. Audit source-level memetakan tujuh komponen kriptografi inti, yaitu "
        "Noise_IK, X25519, ChaCha20-Poly1305, BLAKE2s, Argon2id, Ed25519, dan OsRng. Hasil analisis "
        "menunjukkan bahwa Ed25519 hanya dibangkitkan dan disimpan sebagai bahan fingerprint; tidak "
        "terdapat operasi sign atau verify aktif. Verifikasi identitas bersifat fail-closed untuk "
        "kontak yang sudah dikenal, sedangkan kontak baru mengikuti trust-on-first-use implisit. "
        "Pengujian release menghasilkan 46 dari 46 test lulus tanpa kegagalan. Vault identitas "
        "berukuran tepat 108 byte pada lima sampel, invite tanpa alamat onion konsisten 86 karakter "
        "pada lima keypair, dan sepuluh proses unseal menghasilkan invite identik. Pada laptop "
        "Intel Core i7-1165G7 dengan Windows 11 dan RAM 11,79 GB, tiga puluh pengukuran cold-start "
        "memberikan biaya unseal neto rata-rata 47,99 ms dengan simpangan baku 11,41 ms; hasil ini "
        "mengoreksi estimasi sekitar 100 ms pada komentar kode untuk lingkungan uji tersebut. "
        "Latensi handshake hanya dapat dinyatakan sebagai batas atas kurang dari 0,86 ms, sedangkan "
        "overhead transport Noise dan penggunaan memori puncak belum diukur langsung. Karena itu, "
        "hasil pengujian mendukung correctness dan rejection pada cakupan yang diuji, tetapi tidak "
        "membuktikan keamanan sistem secara absolut maupun memverifikasi forward secrecy secara lokal."
    )
    words = len(re.findall(r"\b[\w-]+\b", abstract))
    if not 200 <= words <= 300:
        raise AssertionError(f"Abstrak {words} kata; harus 200-300")
    add_para(doc, abstract, "Abstract Body")
    add_para(
        doc,
        "Kata kunci: Argon2id; ChaCha20-Poly1305; manajemen kunci; Noise_IK; peer-to-peer; threat model.",
        "Abstract Body",
    )

    doc.add_page_break()
    add_front_heading(doc, "DAFTAR ISI")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    add_field(p, r'TOC \o "1-2" \h \z \u', "Daftar isi diperbarui saat dokumen dibuka.")

    doc.add_page_break()
    add_front_heading(doc, "DAFTAR GAMBAR")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    add_field(p, r'TOC \h \z \t "Caption Figure,1"', "Daftar gambar diperbarui saat dokumen dibuka.")

    doc.add_page_break()
    add_front_heading(doc, "DAFTAR TABEL")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    add_field(p, r'TOC \h \z \t "Caption Table,1"', "Daftar tabel diperbarui saat dokumen dibuka.")


def add_main_section(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(section)
    set_page_numbering(section, "decimal", 1)
    setup_footer(section, roman=False)


def continue_main_page_numbering(doc: Document):
    """Remove inherited restart markers from table-orientation sections."""
    for section in doc.sections[3:]:
        set_page_numbering(section, "decimal")


def build_content(doc: Document, numbering: dict[str, int]):
    diagrams = PREP / "diagrams" / "rendered" / "png"
    screenshots = PREP / "screenshots"

    # BAB I
    add_chapter(
        doc,
        1,
        "Bab ini menjelaskan latar belakang, rumusan masalah, tujuan, manfaat, batasan, "
        "serta sistematika penulisan laporan mini-TA AKSARA.",
    )
    add_subheading(doc, 1, 1, "Latar Belakang")
    add_para(
        doc,
        "Komunikasi digital yang aman memerlukan kerahasiaan, integritas, dan keaslian "
        "tanpa selalu bergantung pada pihak ketiga tepercaya. AKSARA menempatkan kebutuhan "
        "tersebut pada arsitektur *peer-to-peer* dua pihak: setiap pengguna menjalankan "
        "aplikasi terminal, menyimpan identitas secara lokal, dan membentuk sesi langsung "
        "melalui LAN atau Tor. Kanal aman dibangun dengan pola "
        "`Noise_IK_25519_ChaChaPoly_BLAKE2s`, sedangkan ChaCha20-Poly1305 digunakan untuk "
        "perlindungan data *at rest* dan sebagai komponen transport Noise "
        "[[CITE:noise2018,rfc8439]].",
    )
    add_para(
        doc,
        "Objek yang dianalisis adalah implementasi AKSARA v0.2.1 berbasis Rust 2021. "
        "Struktur utamanya meliputi modul `identity`, `crypto`, `transport`, `contacts`, "
        "`session`, dan `tui`, dengan `src/main.rs` sebagai orkestrator. Penelitian ini "
        "tidak merancang protokol kriptografi baru. Fokusnya adalah mendokumentasikan "
        "perilaku *as built*, menelusuri setiap klaim implementasi ke path dan simbol kode, "
        "serta membandingkan hasil pengujian dengan klaim yang sudah ada.",
    )
    add_para(
        doc,
        "Sisi empiris ditempatkan sebagai verifikasi terbatas terhadap *correctness*, "
        "*rejection*, konsistensi format, dan performa. Hasil pengukuran tidak digunakan "
        "untuk menyatakan bahwa AKSARA terbukti aman, melainkan untuk membatasi kesimpulan "
        "pada perilaku yang benar-benar teramati di lingkungan uji.",
    )

    add_subheading(doc, 1, 2, "Rumusan Masalah")
    add_para(
        doc,
        "Berdasarkan latar belakang tersebut, penelitian ini merumuskan tiga permasalahan "
        "utama terkait implementasi protokol, manajemen kunci, dan ketahanan keamanan AKSARA.",
    )
    for index, item in enumerate([
        "Bagaimana AKSARA mengimplementasikan Noise_IK untuk autentikasi identitas dan "
        "kerahasiaan sesi, serta properti mana yang benar-benar dapat diverifikasi dari "
        "evidence source-level?",
        "Bagaimana siklus hidup kunci kriptografi AKSARA dikelola secara end-to-end, dan "
        "sejauh mana pengelolaan tersebut konsisten dengan standar terkait?",
        "Sejauh mana desain protokol dan manajemen kunci AKSARA memberikan ketahanan "
        "terhadap kelas ancaman yang relevan, dan apa batasan atau celah residual yang "
        "teridentifikasi?",
    ]):
        add_lettered_item(doc, item, index)

    add_subheading(doc, 1, 3, "Tujuan Penelitian")
    add_para(
        doc,
        "Sejalan dengan rumusan masalah, penelitian ini memiliki tiga tujuan yang saling "
        "berkaitan sebagai satu rangkaian evaluasi protokol keamanan AKSARA.",
    )
    for index, item in enumerate([
        "Mendokumentasikan dan menganalisis implementasi Noise_IK beserta primitif "
        "kriptografi pendukungnya berdasarkan evidence kode sumber.",
        "Menganalisis pembangkitan, penyimpanan, penggunaan, rotasi, revokasi, dan "
        "pembersihan material kunci sepanjang siklus hidupnya.",
        "Menyusun threat model serta menjalankan pengujian correctness, rejection, "
        "konsistensi, dan performa terbatas yang dapat direplikasi.",
    ]):
        add_lettered_item(doc, item, index)

    add_subheading(doc, 1, 4, "Manfaat Penelitian")
    add_para(
        doc,
        "Secara akademik, penelitian ini memberikan studi kasus implementasi kriptografi "
        "terapan pada aplikasi Rust yang nyata. Struktur audit yang menghubungkan klaim, "
        "path kode, referensi teori, dan data pengujian dapat digunakan sebagai contoh "
        "dokumentasi protokol *as built* pada mata kuliah Implementasi Kriptografi.",
    )
    add_para(
        doc,
        "Secara praktis, penelitian ini menyediakan peta kesenjangan yang dapat ditelusuri, "
        "antara lain ketiadaan rotasi kunci, zeroization yang belum merata pada batas fungsi, "
        "dan ketergantungan pada verifikasi fingerprint manual. Temuan tersebut bukan "
        "perbaikan yang sudah diimplementasikan, melainkan dasar prioritas pengembangan.",
    )
    add_para(
        doc,
        "Pembagian kerja kelompok mengikuti peran aktual. Tabel 1.1 memperlihatkan bahwa "
        "implementasi, dokumentasi, dan pengujian menghasilkan luaran yang berbeda namun "
        "saling melengkapi; pembagian ini tidak mengarang kepemilikan modul per anggota.",
    )
    add_table_from_file(
        doc,
        "Tabel 1.1 Pembagian Tugas Anggota",
        "13_team_assignment.md",
        landscape=True,
    )

    add_subheading(doc, 1, 5, "Batasan Masalah")
    add_para(
        doc,
        "Untuk menjaga penelitian tetap proporsional dengan cakupan tugas mata kuliah, "
        "batasan berikut diterapkan sejak awal.",
    )
    limitations = [
        "Keamanan internal crate pihak ketiga, termasuk `snow`, `arti-client`, "
        "`mdns-sd`, dan keluarga `dalek`, tidak diaudit; yang dinilai hanya cara AKSARA "
        "memakainya.",
        "Kebenaran CSPRNG tingkat sistem operasi di balik `OsRng` tidak dinilai.",
        "Penelitian tidak melakukan pembuktian formal kriptografi.",
        "Analisis side-channel fisik dan serangan tingkat perangkat keras tidak dilakukan.",
        "Ketahanan denial-of-service tidak diukur secara kuantitatif pada skala luas.",
        "Skenario multi-device dan sinkronisasi kunci tidak termasuk objek penelitian.",
        "Penelitian bersifat deskriptif-evaluatif dan tidak mengimplementasikan remediasi.",
        "Transport LAN dan Tor dibahas sebagai konteks pembawa sesi, bukan objek evaluasi "
        "kriptografi primer; mDNS hanya didukung evidence source-level dan Tor oleh bukti "
        "visual, bukan eksperimen terukur penuh.",
        "Klaim performa dibatasi pada pengukuran aktual di satu unit hardware dan tidak "
        "digeneralisasi ke perangkat lain.",
    ]
    for index, item in enumerate(limitations):
        add_lettered_item(doc, item, index)

    add_subheading(doc, 1, 6, "Sistematika Penulisan")
    add_para(
        doc,
        "Laporan disusun dalam enam bab. BAB I memuat pendahuluan. BAB II membahas teori "
        "kriptografi dan penelitian terkait. BAB III menjelaskan metode audit dan "
        "pengujian. BAB IV menguraikan arsitektur, implementasi protokol, siklus hidup "
        "kunci, serta threat model. BAB V menyajikan hasil pengujian dan analisisnya. "
        "BAB VI menjawab rumusan masalah, menyatakan keterbatasan, dan memberi saran.",
    )

    # BAB II
    add_chapter(
        doc,
        2,
        "Bab ini membangun landasan teori untuk menilai implementasi AKSARA. Pembahasan "
        "membedakan properti algoritma dari bukti implementasi lokal agar klaim keamanan "
        "tidak melampaui evidence.",
    )
    add_subheading(doc, 2, 1, "Konsep Dasar Keamanan Informasi")
    add_para(
        doc,
        "Kerahasiaan membatasi isi pesan kepada pihak yang berhak; integritas memungkinkan "
        "perubahan terdeteksi; dan keaslian mengikat komunikasi kepada identitas yang "
        "diharapkan. Dalam AKSARA, dua properti pertama terutama diberikan oleh AEAD, "
        "sedangkan keaslian peer berasal dari static key Noise_IK dan hanya bersifat "
        "fail-closed bagi kontak yang sudah dikenal [[CITE:rfc8439,noise2018]].",
    )
    add_para(
        doc,
        "Pemetaan pada Tabel 2.1 menunjukkan bahwa kerahasiaan dan integritas mempunyai "
        "evidence kuat pada vault serta contact store, tetapi autentikasi dan zeroization "
        "masih parsial. Status tersebut merupakan hasil audit, bukan skor keamanan.",
    )
    add_table_from_file(
        doc,
        "Tabel 2.1 Kebutuhan Non-Fungsional dan Tingkat Pemenuhan",
        "02_nonfunctional_requirements.md",
        landscape=True,
    )

    add_subheading(doc, 2, 2, "AEAD (Authenticated Encryption with Associated Data)")
    add_para(
        doc,
        "*Authenticated Encryption with Associated Data* menggabungkan enkripsi dan tag "
        "autentikasi dalam satu antarmuka. ChaCha20 membentuk *keystream*, sedangkan "
        "Poly1305 menghasilkan tag 128 bit; konstruksi IETF memakai nonce 96 bit "
        "[[CITE:rfc8439,bernstein2008chacha,bernstein2005poly1305]].",
    )
    add_para(
        doc,
        "AKSARA menggunakan ChaCha20-Poly1305 pada vault (`identity::vault`), contact "
        "store (`contacts::save_contacts`/`load_contacts`), dan transport melalui suite "
        "Noise. Ketiga konteks tidak memakai AAD. Algoritma ini juga tidak "
        "*misuse-resistant*: keunikan nonce tetap menjadi prasyarat, sehingga keberadaan "
        "nonce acak tidak boleh ditafsirkan sebagai perlindungan terhadap nonce reuse.",
    )
    add_para(
        doc,
        "Implementasi at-rest menggunakan crate `chacha20poly1305` 0.10.1 "
        "[[CITE:chacha20poly1305crate]]. Untuk konteks transport, pemilihan ChaChaPoly "
        "berasal dari nama suite Noise dan dikelola internal oleh `snow`; karena itu "
        "confidence-nya lebih rendah daripada dua pemanggilan langsung di aplikasi.",
    )

    add_subheading(doc, 2, 3, "MAC dan Fungsi Hash")
    add_para(
        doc,
        "MAC memberikan autentikasi dan integritas pesan dengan kunci rahasia. Pada AKSARA, "
        "Poly1305 tidak dipanggil sebagai MAC berdiri sendiri, melainkan sebagai komponen "
        "ChaCha20-Poly1305. BLAKE2s mempunyai peran berbeda: fingerprint identitas, "
        "derivasi contacts-key secara single-shot, dan hash internal suite Noise "
        "[[CITE:rfc7693,aumasson2013blake2]].",
    )
    add_para(
        doc,
        "Fingerprint di `contacts::fingerprint` memasukkan public key Ed25519 dan X25519 "
        "dengan context string `aksara-fingerprint-v1`. Derivasi contacts-key memakai "
        "context string lain, `aksara-contacts-key-v1`, sehingga kedua domain dibedakan. "
        "Namun hashing fingerprint tidak sama dengan integritas pesan; integritas pesan "
        "tetap berasal dari AEAD.",
    )
    add_para(
        doc,
        "SHA-256 dan SHA3-256 merupakan alternatif standar yang sah "
        "[[CITE:fips180-4,fips202]], tetapi BLAKE2s dipertahankan karena sudah menjadi "
        "komponen suite Noise dan tersedia melalui crate `blake2` "
        "[[CITE:blake2crate]].",
    )

    add_subheading(doc, 2, 4, "Key Agreement (Diffie-Hellman/ECDH)")
    add_para(
        doc,
        "*Key agreement* memungkinkan dua pihak menurunkan rahasia bersama melalui kanal "
        "publik. X25519 merupakan fungsi Diffie-Hellman berbasis Curve25519 dengan input "
        "dan output 32 byte [[CITE:rfc7748,bernstein2006curve25519]].",
    )
    add_para(
        doc,
        "Pola IK memakai empat token DH, yaitu `es`, `ss`, `ee`, dan `se`. Static key "
        "dibangkitkan melalui `NoiseKey::generate`, sedangkan ephemeral key dikelola "
        "internal oleh `snow`. Operasi DH aktual tidak terlihat sebagai pemanggilan "
        "langsung pada source aplikasi, sehingga audit membedakan confidence keygen yang "
        "tinggi dari orkestrasi internal yang medium.",
    )
    add_para(
        doc,
        "P-256 dan Curve448 merupakan alternatif dengan parameter standar "
        "[[CITE:sp800-186,fips186-5]], tetapi X25519 dipilih karena sesuai token suite "
        "Noise, bentuk antarmuka ringkas, dan dukungan crate `x25519-dalek` "
        "[[CITE:x25519dalekcrate]].",
    )

    add_subheading(doc, 2, 5, "Key Derivation Function (KDF)")
    add_para(
        doc,
        "KDF mengubah material masukan menjadi kunci dengan panjang dan distribusi yang "
        "sesuai. Passphrase manusia berentropi rendah memerlukan KDF yang mahal terhadap "
        "serangan paralel. AKSARA memakai Argon2id dengan `m=19 MiB`, `t=2`, `p=1`, dan "
        "output 32 byte untuk vault [[CITE:rfc9106,biryukov2016argon2]].",
    )
    add_para(
        doc,
        "Alternatif scrypt juga memory-hard [[CITE:rfc7914,percival2009scrypt]], sedangkan "
        "PBKDF2 tetap bergantung terutama pada iterasi komputasi dan bukan memori "
        "[[CITE:rfc8018]]. Pemilihan Argon2id diimplementasikan oleh crate `argon2` 0.5.3 "
        "[[CITE:argon2crate]].",
    )
    add_para(
        doc,
        "Contacts-key tidak berasal dari passphrase, melainkan dari secret key identitas "
        "yang berentropi tinggi. `derive_contacts_key` memakai BLAKE2s single-shot dengan "
        "domain separation. Konstruksi ini bukan HKDF standar dan dicatat sebagai keputusan "
        "desain dengan cakupan pemakaian tunggal.",
    )

    add_subheading(doc, 2, 6, "Identitas Digital dan Digital Signature")
    add_para(
        doc,
        "Ed25519 adalah instansiasi EdDSA yang mendukung tanda tangan digital deterministik "
        "dengan kunci publik ringkas [[CITE:rfc8032,bernstein2012ed25519,fips186-5]]. "
        "Kemampuan algoritma tersebut tidak otomatis menjadi kemampuan aplikasi.",
    )
    add_para(
        doc,
        "Pada AKSARA, `IdentityKey::generate` membangkitkan keypair Ed25519 dan public key "
        "menjadi bahan fingerprint. Audit menyeluruh tidak menemukan pemanggilan "
        "`sign()`, `verify()`, atau tipe `Signature`. Dengan demikian, AKSARA tidak "
        "menggunakan tanda tangan digital Ed25519 secara aktif; identitas operasional "
        "diverifikasi melalui fingerprint dan static key Noise_IK. Crate "
        "`ed25519-dalek` menyediakan API yang lebih luas, tetapi API signing tersebut "
        "belum digunakan [[CITE:ed25519dalekcrate]].",
    )

    add_subheading(doc, 2, 7, "Nonce dan Replay Protection")
    add_para(
        doc,
        "Nonce pada ChaCha20-Poly1305 harus unik untuk setiap kunci. Vault dan contact "
        "store membangkitkan nonce 96 bit baru dari `OsRng` pada setiap operasi. Nonce "
        "transport dikelola internal oleh `snow`, sehingga tidak dapat diaudit langsung "
        "dari `session::run_session` [[CITE:rfc8439]].",
    )
    add_para(
        doc,
        "Replay protection biasanya membutuhkan nomor urut, jendela penerimaan, atau "
        "mekanisme protokol setara. Tidak ditemukan sequence number atau pemeriksaan "
        "timestamp eksplisit pada payload sesi AKSARA. Oleh karena itu laporan tidak "
        "menyatakan replay protection tersedia hanya karena protokol memakai nonce.",
    )

    add_subheading(doc, 2, 8, "Manajemen Kunci")
    add_para(
        doc,
        "Siklus hidup kunci mencakup pembangkitan, penyimpanan, penggunaan, rotasi, "
        "revokasi, dan penghancuran. AKSARA memakai `rand::rngs::OsRng` pada pembangkitan "
        "kunci, salt, dan nonce; kualitas implementasi CSPRNG tingkat OS berada di luar "
        "cakupan penelitian [[CITE:sp800-90a,randcrate]].",
    )
    add_para(
        doc,
        "Material jangka panjang disimpan di vault terenkripsi. Tipe inti memakai "
        "`ZeroizeOnDrop` atau `Zeroizing`, tetapi boundary fungsi tertentu masih "
        "menyalin secret ke buffer biasa. Audit juga tidak menemukan rotasi atau revokasi "
        "in-band. Karena itu, properti forward secrecy yang didokumentasikan oleh pola "
        "Noise_IK tetap berstatus `DOCUMENTED_ONLY`, bukan hasil verifikasi lokal.",
    )

    add_subheading(doc, 2, 9, "Noise Protocol Framework dan Pola Handshake IK")
    add_para(
        doc,
        "Noise Protocol Framework menyusun protokol handshake dari pola pesan dan pilihan "
        "primitif. Pola IK mengasumsikan initiator telah mengetahui static key responder. "
        "Dua pesan polanya adalah `-> e, es, s, ss` dan `<- e, ee, se` "
        "[[CITE:noise2018]].",
    )
    add_para(
        doc,
        "AKSARA menggunakan string literal "
        "`Noise_IK_25519_ChaChaPoly_BLAKE2s` di `crypto::handshake`. Crate `snow` 0.10.0 "
        "mengelola state handshake dan transport; dokumentasi crate menyatakan implementasi "
        "belum diaudit formal [[CITE:snowcrate]]. Hal ini menjadi batas kepercayaan penting: "
        "orkestrasi aplikasi dapat diaudit, sedangkan mekanisme internal crate tidak.",
    )
    add_para(
        doc,
        "Tabel 2.2 merangkum alasan pemilihan tujuh komponen inti dan keterbatasannya. "
        "Tabel 2.3 menempatkan pilihan itu terhadap alternatif yang valid; penolakan "
        "alternatif merupakan keputusan kesesuaian konteks, bukan penilaian bahwa alternatif "
        "tersebut tidak aman.",
    )
    add_table_from_file(
        doc,
        "Tabel 2.2 Justifikasi Algoritma Kriptografi",
        "05_algorithm_justification.md",
        landscape=True,
    )
    add_table_from_file(
        doc,
        "Tabel 2.3 Perbandingan Alternatif Algoritma",
        "06_algorithm_alternative_comparison.md",
        landscape=True,
    )

    add_subheading(doc, 2, 10, "Penelitian Terkait")
    add_para(
        doc,
        "Posisi AKSARA dibandingkan melalui tiga sumbu: kerangka kriptografi, siklus hidup "
        "kunci, dan arsitektur jaringan. Noise Explorer menunjukkan bahwa pola Noise dapat "
        "dimodelkan secara formal [[CITE:kobeissi2019noiseexplorer]], sedangkan analisis "
        "Signal dan Matrix memperlihatkan manfaat ratcheting serta rotasi sesi "
        "[[CITE:cohngordon2020signal,albrecht2024matrix]].",
    )
    add_para(
        doc,
        "WireGuard memakai stack Noise IK, X25519, ChaCha20-Poly1305, dan BLAKE2s yang dekat "
        "dengan AKSARA, tetapi mempunyai mekanisme rekey dan anti-DoS yang berbeda "
        "[[CITE:donenfeld2017wireguard]]. OTR menekankan deniability dan forward secrecy "
        "[[CITE:borisov2004otr]], sedangkan Briar dan Tox memberikan pembanding P2P dengan "
        "jalur serta mekanisme penemuan yang berbeda [[CITE:briarspec,toxspec]].",
    )
    add_para(
        doc,
        "Alternatif protokol yang turut dipertimbangkan adalah TLS 1.3 serta X3DH. TLS "
        "mempunyai model sertifikat/PKI yang tidak natural bagi identitas self-sovereign "
        "AKSARA, sedangkan X3DH dirancang untuk pembentukan sesi asinkron yang biasanya "
        "menggunakan server perantara [[CITE:rfc8446,x3dh2016]]. Pada tingkat AEAD, "
        "AES-GCM dan AES-GCM-SIV adalah alternatif standar, termasuk pilihan yang lebih "
        "toleran terhadap misuse nonce [[CITE:sp800-38d,rfc8452]].",
    )
    add_para(
        doc,
        "Tabel 2.4 menunjukkan lima gap yang dirumuskan sebagai “belum ditemukan pada "
        "sumber yang ditinjau”, bukan klaim universal bahwa topik tersebut belum pernah "
        "diteliti. Gap utama adalah ketiadaan verifikasi formal instansiasi AKSARA, "
        "ratcheting, evaluasi overhead jaringan, handshake kontak terstruktur, dan fallback "
        "offline.",
    )
    add_table_from_file(
        doc,
        "Tabel 2.4 Penelitian Terkait",
        "10_related_work.md",
        landscape=True,
    )

    # BAB III
    add_chapter(
        doc,
        3,
        "Bab ini menjelaskan rancangan studi kasus, objek, tahapan audit, lingkungan, dan "
        "skenario pengujian. Hasil tidak dicampurkan ke bab metodologi.",
    )
    add_subheading(doc, 3, 1, "Jenis Penelitian")
    add_para(
        doc,
        "Penelitian ini merupakan studi kasus deskriptif-evaluatif terhadap implementasi "
        "kriptografi nyata. Pendekatan utamanya adalah audit source-level yang memetakan "
        "klaim ke simbol kode, dilengkapi evaluasi threat model serta eksperimen "
        "*correctness*, *rejection*, konsistensi, dan performa terbatas.",
    )
    add_para(
        doc,
        "Kajian literatur dilakukan secara terarah untuk memvalidasi teori primitif dan "
        "membandingkan sistem sejenis; kajian tersebut bukan systematic literature review "
        "formal. Kesimpulan dibatasi pada evidence kode dan data yang benar-benar tersedia.",
    )

    add_subheading(doc, 3, 2, "Objek Penelitian")
    add_para(
        doc,
        "Objek penelitian adalah AKSARA v0.2.1, sebuah binary tunggal Rust dengan entry "
        "point `src/main.rs`. Analisis berfokus pada `identity`, `crypto`, `transport`, "
        "`contacts`, `session`, dan `tui`. Commit implementasi yang terdokumentasi pada "
        "audit adalah `450d484`; status build tersebut merupakan catatan sesi sebelumnya, "
        "bukan verifikasi baru pada penyusunan laporan ini.",
    )
    add_para(
        doc,
        "Unit analisis kriptografi dikonsolidasikan menjadi CORE-1 sampai CORE-7. Transport "
        "LAN/Tor hanya dinilai sejauh memengaruhi pembentukan sesi, metadata, dan trust "
        "boundary.",
    )

    add_subheading(doc, 3, 3, "Tahapan Penelitian")
    add_para(
        doc,
        "Penelitian dilaksanakan melalui lima tahapan berurutan agar temuan dapat ditelusuri "
        "dan tidak bergantung pada dokumentasi saja.",
    )
    for index, item in enumerate([
        "Mengaudit struktur codebase dan mengumpulkan klaim source-level per modul.",
        "Menormalisasi 36 entry kriptografi menjadi tujuh komponen inti, lalu menyusun "
        "justifikasi dan perbandingan alternatif.",
        "Mendokumentasikan spesifikasi protokol, siklus hidup kunci, dan threat model "
        "berdasarkan perilaku as-built.",
        "Menyusun skenario serta metrik pengujian untuk correctness, rejection, konsistensi, "
        "dan performa.",
        "Menjalankan pengujian, mencatat lingkungan dan data mentah, kemudian membandingkan "
        "hasil dengan klaim implementasi tanpa menaikkan status klaim yang tidak teruji.",
    ]):
        add_lettered_item(doc, item, index)

    add_subheading(doc, 3, 4, "Lingkungan dan Alat")
    add_para(
        doc,
        "Analisis menggunakan toolchain Rust, pencarian kode terarah, pembacaan simbol, dan "
        "matriks evidence. Diagram teknis dirender dengan Mermaid CLI. Pengujian aktual "
        "dijalankan dengan `rustc` dan `cargo` 1.97.0 pada profil `--release`; spesifikasi "
        "hardware lengkap disajikan pada BAB V agar hasil performa tidak dipisahkan dari "
        "lingkungannya.",
    )
    add_para(
        doc,
        "Tabel 3.1 merangkum dependency utama berdasarkan `Cargo.lock`. Versi ini digunakan "
        "sebagai ground truth; dokumentasi crate hanya menjelaskan API dan tidak menjadi "
        "bukti bahwa suatu fitur benar-benar dipanggil aplikasi.",
    )
    add_table_from_file(
        doc,
        "Tabel 3.1 Stack Teknologi AKSARA",
        "03_tech_stack.md",
        landscape=True,
    )

    add_subheading(doc, 3, 5, "Rencana Pengujian")
    add_para(
        doc,
        "Lima kelompok eksperimen dirancang sebelum hasil ditafsirkan. EXP-01 menguji vault; "
        "EXP-02 menguji handshake; EXP-03 menguji transport dan framing; EXP-04 menguji "
        "invite, fingerprint, dan contact store; EXP-05 mengukur biaya Argon2id dan ukuran "
        "vault. Metrik, hasil, dan keterbatasan pengukuran disajikan pada BAB V.",
    )
    add_para(
        doc,
        "Known-answer test tidak dimasukkan karena test suite tidak menggunakan test vector "
        "standar. Modified-AAD rejection tidak berlaku karena AKSARA tidak memakai AAD, "
        "sedangkan replay rejection tidak dapat diuji karena mekanismenya tidak ditemukan. "
        "Ketiadaan kandidat tersebut dicatat, bukan diganti dengan hasil buatan.",
    )

    # BAB IV
    add_chapter(
        doc,
        4,
        "Bab inti ini menjelaskan arsitektur dan implementasi AKSARA secara rinci. Setiap "
        "klaim implementasi dihubungkan ke path atau simbol kode, sedangkan gambar dan tabel "
        "menjadi ringkasan visual atas evidence tersebut.",
    )
    add_subheading(doc, 4, 1, "Gambaran Umum Arsitektur AKSARA")
    add_para(
        doc,
        "AKSARA menjalankan dua proses setara tanpa broker pusat. `transport::role_from_fp` "
        "menentukan peran initiator atau responder secara deterministik dari perbandingan "
        "fingerprint, bukan konfigurasi client/server tetap. Gambar 4.1 memperlihatkan trust "
        "boundary antara pengguna, proses, jaringan, dan penyimpanan lokal.",
    )
    add_figure(
        doc,
        diagrams / "01-context.png",
        "Gambar 4.1 Diagram Konteks AKSARA",
    )
    add_para(
        doc,
        "Secara modular, `src/main.rs` mengorkestrasi identitas, kontak, transport, dan TUI; "
        "`transport` menyediakan stream, `crypto::handshake` mengubahnya menjadi state "
        "Noise, dan `session::run_session` menangani pesan. Relasi ini dirangkum pada "
        "Gambar 4.2.",
    )
    add_figure(
        doc,
        diagrams / "02-component-architecture.png",
        "Gambar 4.2 Diagram Arsitektur Komponen",
    )
    add_para(
        doc,
        "Tampilan aktual pada Gambar 4.3 memperlihatkan TUI dengan badge `⌂ LOCAL`, "
        "fingerprint pendek, panel kontak, dan bantuan tombol. Badge tersebut menunjukkan "
        "status Tor tidak aktif; ia bukan label bahwa sesi tertentu telah memakai LAN.",
    )
    add_figure(
        doc,
        screenshots / "01-antarmuka-utama.png",
        "Gambar 4.3 Antarmuka Utama AKSARA",
    )
    add_para(
        doc,
        "Alur end-to-end bergerak dari pembukaan vault, pertukaran invite out-of-band, "
        "pembentukan transport, handshake, hingga sesi terenkripsi. Gambar 4.4 menempatkan "
        "urutan tersebut sebagai konteks bagi subbab berikutnya.",
    )
    add_figure(
        doc,
        diagrams / "04-sequence-main-flow.png",
        "Gambar 4.4 Sequence Diagram Proses Utama",
    )
    add_para(
        doc,
        "Tabel 4.1 merangkum kebutuhan fungsional yang ditemukan pada source. Status "
        "`IMPLEMENTED` menunjukkan keberadaan path dan simbol, bukan jaminan bahwa seluruh "
        "skenario keamanan telah diuji.",
    )
    add_table_from_file(
        doc,
        "Tabel 4.1 Kebutuhan Fungsional AKSARA",
        "01_functional_requirements.md",
        landscape=True,
    )

    add_subheading(doc, 4, 2, "Identitas dan Manajemen Kunci")
    add_para(
        doc,
        "`IdentityKey::generate` dan `NoiseKey::generate` memakai `OsRng` untuk keypair "
        "Ed25519 dan X25519 yang terpisah. Pemisahan ini menghindari penggunaan satu kunci "
        "untuk tujuan identitas dan DH sekaligus. `KeyBundle` membungkus kedua secret dan "
        "menerapkan `ZeroizeOnDrop`.",
    )
    add_para(
        doc,
        "Vault mempunyai layout tetap 108 byte: salt 16 byte, nonce 12 byte, ciphertext "
        "64 byte, dan tag 16 byte. `identity::vault::seal` menurunkan kunci 32 byte dengan "
        "Argon2id lalu mengenkripsi 64 byte secret. `unseal` memetakan passphrase salah, "
        "ukuran tidak valid, dan tag gagal ke pesan generik `vault could not be opened`, "
        "sehingga perilakunya fail-closed dan tidak membedakan penyebab kepada pengguna.",
    )
    add_para(
        doc,
        "Contact store memakai contacts-key yang diturunkan dari identity secret melalui "
        "BLAKE2s dan nonce acak pada setiap `save_contacts`. Secret inti dan buffer vault "
        "dibersihkan, tetapi audit menemukan boundary biasa pada `session`, `handshake`, "
        "`contacts`, dan `main.rs`. Input passphrase stdin juga masih terlihat di layar.",
    )
    add_para(
        doc,
        "Tidak ditemukan rotasi atau revokasi untuk identity key, Noise static key, maupun "
        "parameter vault. Gambar 4.5 menampilkan siklus hidup aktual tanpa menambahkan "
        "transisi yang belum diimplementasikan.",
    )
    add_figure(
        doc,
        diagrams / "06-key-lifecycle-state.png",
        "Gambar 4.5 Diagram Siklus Hidup Kunci",
    )
    add_para(
        doc,
        "Tabel 4.2 memetakan material kunci, ukuran, umur, dan persistensinya. Status "
        "zeroization dan ketiadaan rotasi dijelaskan dalam narasi karena merupakan sifat "
        "lintas material.",
    )
    add_table_from_file(
        doc,
        "Tabel 4.2 Lifecycle Kunci",
        "08_key_lifecycle.md",
        landscape=True,
    )

    add_subheading(doc, 4, 3, "Spesifikasi Protokol: Invite, Discovery, dan Pembentukan Koneksi")
    add_para(
        doc,
        "`contacts::encode_invite` menggabungkan public key Ed25519 32 byte dan X25519 "
        "32 byte, lalu melakukan base64url tanpa padding. Bentuk LAN-only mempunyai panjang "
        "86 karakter; suffix `@<onion-address>` bersifat opsional. Invite tidak mempunyai "
        "tanda tangan atau autentikasi kriptografis tersendiri, sehingga verifikasi "
        "fingerprint out-of-band tetap penting.",
    )
    add_para(
        doc,
        "Gambar 4.6 memperlihatkan panel identitas aktual dengan invite 86 karakter dan "
        "fingerprint 64 heksadesimal. Identitas pada gambar adalah data dummy.",
    )
    add_figure(
        doc,
        screenshots / "02-identitas-invite.png",
        "Gambar 4.6 Panel Identitas dan Invite Code",
    )
    add_para(
        doc,
        "Discovery LAN melalui `transport::lan::advertise` mengiklankan service "
        "`_aksara._tcp.local.` dan fingerprint melalui TXT record. X25519 public key tidak "
        "diiklankan, tetapi presence dan fingerprint tetap terlihat plaintext sesuai sifat "
        "mDNS/DNS-SD [[CITE:rfc6762,rfc6763]]. Jalur advertise/browse ini mempunyai "
        "evidence source-level dan belum diverifikasi eksperimen resmi.",
    )
    add_para(
        doc,
        "`transport::establish` mencoba LAN terlebih dahulu. Timeout tiga detik hanya "
        "diterapkan pada mode Auto ketika Tor tersedia; setelah itu alur beralih ke Tor. "
        "Dial Tor mengulang setiap delapan detik hingga total 120 detik. Gambar 4.7 "
        "memperlihatkan invite dengan alamat onion; ia membuktikan onion address telah "
        "dibentuk, bukan mengukur performa jalur Tor.",
    )
    add_figure(
        doc,
        screenshots / "05-onion-invite.png",
        "Gambar 4.7 Invite dengan Onion Address",
    )

    add_subheading(doc, 4, 4, "Handshake Noise_IK")
    add_para(
        doc,
        "`crypto::handshake` membangun state dari string literal "
        "`Noise_IK_25519_ChaChaPoly_BLAKE2s`. Initiator menulis pesan "
        "`e, es, s, ss`, responder menjawab `e, ee, se`, kemudian keduanya berpindah ke "
        "transport mode. Urutan dua pesan teruji oleh `handshake_ik_roundtrip`.",
    )
    add_para(
        doc,
        "Pada `session::run_session`, initiator harus memiliki `peer_noise_pk`; bila tidak, "
        "sesi gagal sebelum handshake. Responder memanggil `remote_static()` dan "
        "membandingkannya dengan key yang diharapkan hanya jika kontak sudah dikenal. "
        "Ketidakcocokan menghasilkan `IdentityMismatch`.",
    )
    add_para(
        doc,
        "Jika kontak belum dikenal dan `peer_noise_pk=None`, blok perbandingan tidak "
        "dijalankan. Fakta ketiadaan pemeriksaan ini berconfidence tinggi, sedangkan "
        "interpretasi bahwa trust-on-first-use tersebut disengaja masih memerlukan "
        "konfirmasi desain. Gambar 4.8 memvisualisasikan percabangan itu.",
    )
    add_figure(
        doc,
        diagrams / "05-sequence-handshake-noise-ik.png",
        "Gambar 4.8 Sequence Diagram Handshake Noise_IK",
    )
    add_para(
        doc,
        "Forward secrecy, kerahasiaan static key initiator terhadap penyadap pasif, dan "
        "mutual authentication merupakan properti yang didokumentasikan pada pola IK, tetapi "
        "test AKSARA tidak memverifikasinya secara penuh. Laporan mempertahankan status "
        "`DOCUMENTED_ONLY` atau `PARTIAL` dan tidak menyatakannya terbukti.",
    )

    add_subheading(doc, 4, 5, "Transport Sesi Terenkripsi")
    add_para(
        doc,
        "Setelah handshake, `session::run_session` memakai tag plaintext satu byte: "
        "`TYPE_TEXT=0x00`, `TYPE_BLUR=0x01`, dan `TYPE_PING=0x02`. Payload kemudian "
        "dienkripsi oleh state transport Noise dan dibungkus oleh frame dengan prefix panjang "
        "dua byte big-endian. Batas `MAX_FRAME_LEN` adalah 65.535 byte.",
    )
    add_para(
        doc,
        "Model konkurensi menggunakan task pembaca terpisah karena `read_frame` tidak "
        "cancel-safe. Pesan oversize menghasilkan notice non-fatal, sedangkan kegagalan "
        "dekripsi membuat loop berhenti secara fail-closed. UI belum membedakan penutupan "
        "akibat tampering dari penutupan normal.",
    )
    add_para(
        doc,
        "Gambar 4.9 dan Gambar 4.10 menampilkan dua sisi sesi uji dengan topologi TCP "
        "loopback lokal. Kedua gambar membuktikan pertukaran pesan pada aplikasi, tetapi "
        "tidak boleh ditafsirkan sebagai bukti LAN fisik atau Tor.",
    )
    add_figure(
        doc,
        screenshots / "03-komunikasi-loopback-a.png",
        "Gambar 4.9 Sesi Terenkripsi — Sisi Inisiator (Topologi Loopback Lokal)",
    )
    add_figure(
        doc,
        screenshots / "03-komunikasi-loopback-b.png",
        "Gambar 4.10 Sesi Terenkripsi — Sisi Responder (Topologi Loopback Lokal)",
    )
    add_para(
        doc,
        "Sesi lintas jaringan pada Gambar 4.11 dan Gambar 4.12 dibaca bersama bukti subnet "
        "Gambar 4.13 dan Gambar 4.14. Laptop 1 berada pada `192.168.102.128/24`, sedangkan "
        "Laptop 2 pada `192.168.93.113/22`. Karena `establish()` hanya mempunyai jalur LAN "
        "lokal dan fallback Tor setelah `LAN_AUTO_TIMEOUT`, kombinasi sesi aktif dan subnet "
        "berbeda mendukung kesimpulan bahwa jalur yang dipakai adalah Tor. Chat screenshot "
        "sendirian tidak cukup untuk kesimpulan tersebut.",
    )
    add_figure(
        doc,
        screenshots / "07a-komunikasi-tor-laptop1.png",
        "Gambar 4.11 Sesi Terenkripsi Lintas Jaringan via Tor — Sisi Laptop 1",
    )
    add_figure(
        doc,
        screenshots / "07b-komunikasi-tor-laptop2.png",
        "Gambar 4.12 Sesi Terenkripsi Lintas Jaringan via Tor — Sisi Laptop 2",
    )
    add_figure(
        doc,
        screenshots / "07c-bukti-jaringan-berbeda-laptop1.png",
        "Gambar 4.13 Bukti Jaringan Berbeda — Laptop 1",
    )
    add_figure(
        doc,
        screenshots / "07c-bukti-jaringan-berbeda-laptop2.png",
        "Gambar 4.14 Bukti Jaringan Berbeda — Laptop 2",
    )
    add_para(
        doc,
        "Gambar 4.15 merangkum format biner, sedangkan Tabel 4.3 menyajikan field dan "
        "ukurannya. Overhead tag 16 byte pada transport berasal dari spesifikasi suite; "
        "pengukuran langsung pada instance Noise belum dilakukan.",
    )
    add_figure(
        doc,
        diagrams / "07-packet-format.png",
        "Gambar 4.15 Diagram Format Paket/Pesan",
    )
    add_table(
        doc,
        "Tabel 4.3 Format Paket dan Data",
        custom_packet_table(),
        landscape=True,
    )

    add_subheading(doc, 4, 6, "Implementasi Primitif Kriptografi")
    add_para(
        doc,
        "Tujuh komponen inti membentuk rangkaian dari entropi hingga transport, sebagaimana "
        "dirangkum Gambar 4.16. Pemanggilan langsung mempunyai confidence lebih tinggi "
        "daripada perilaku yang hanya diinferensikan dari suite Noise.",
    )
    add_figure(
        doc,
        diagrams / "03-crypto-architecture.png",
        "Gambar 4.16 Diagram Arsitektur Kriptografi",
    )
    add_para(
        doc,
        "CORE-1 memakai `snow`; CORE-2 memakai `x25519-dalek`; CORE-3 memakai "
        "`chacha20poly1305`; CORE-4 memakai `blake2`; CORE-5 memakai `argon2`; CORE-6 "
        "memakai `ed25519-dalek`; dan CORE-7 memakai `rand::OsRng`. Dokumentasi library "
        "menjelaskan implementasi API masing-masing "
        "[[CITE:snowcrate,x25519dalekcrate,chacha20poly1305crate,blake2crate,argon2crate,ed25519dalekcrate,randcrate]], "
        "tetapi bukti pemakaian AKSARA tetap berasal dari source dan `Cargo.lock`.",
    )
    add_para(
        doc,
        "Gambar 4.17 menunjukkan badge `◉ ONLINE`, yang menandakan Tor telah aktif pada "
        "sisi laptop 1. Notifikasi “Tor siap” tidak terlihat pada capture dan karenanya tidak "
        "dikutip sebagai bukti.",
    )
    add_figure(
        doc,
        screenshots / "06-tor-online.png",
        "Gambar 4.17 Badge Transport ◉ ONLINE",
    )
    add_para(
        doc,
        "Tabel 4.4 mempertahankan confidence per komponen. Ed25519 berstatus parsial karena "
        "keygen tersedia tetapi sign/verify tidak dipakai; submekanisme Noise internal juga "
        "tidak diratakan menjadi confidence tinggi.",
    )
    add_table_from_file(
        doc,
        "Tabel 4.4 Inventarisasi Primitif Kriptografi",
        "04_crypto_primitives_inventory.md",
        landscape=True,
    )

    add_subheading(doc, 4, 7, "Threat Model dan Analisis Risiko")
    add_para(
        doc,
        "Trust boundary mencakup jaringan tidak tepercaya, filesystem lokal, proses "
        "aplikasi, interaksi manusia, dan dependency eksternal. Lima model musuh meliputi "
        "penyadap pasif, penyerang LAN aktif, pembaca filesystem, kontak berbahaya yang sah, "
        "dan partisipan mDNS berbahaya.",
    )
    add_para(
        doc,
        "Risk register T1–T7 pada Tabel 4.5 bersifat kualitatif. T1 dan T3 berhubungan "
        "dengan kontak baru dan invite tidak diautentikasi; T2 dengan metadata mDNS; T4 dan "
        "T6 dengan permission lokal/Tor; T5 dengan sinyal UI; serta T7 dengan ketiadaan "
        "rotasi dan revokasi.",
    )
    add_table_from_file(
        doc,
        "Tabel 4.5 Threat Model dan Risk Register",
        "09_threat_model.md",
        landscape=True,
    )
    add_para(
        doc,
        "Mitigasi yang sudah ada dicatat apa adanya, misalnya fingerprint manual, enkripsi "
        "vault, dan fail-closed decryption. Tidak ada skor CVSS atau klaim aman absolut "
        "karena keduanya tidak didukung metode penelitian.",
    )

    # BAB V
    add_chapter(
        doc,
        5,
        "Bab ini menyajikan lingkungan, hasil, dan analisis pengujian. Angka performa selalu "
        "dibaca bersama ukuran sampel, sebaran, dan konfigurasi hardware.",
    )
    add_subheading(doc, 5, 1, "Lingkungan Pengujian Aktual")
    add_para(
        doc,
        "Seluruh pengujian dijalankan pada laptop bare-metal LENOVO 82FG dengan Intel Core "
        "i7-1165G7 4 core/8 thread pada 2,80 GHz, RAM total 11,79 GB, dan sekitar 0,93 GB "
        "bebas saat pengukuran. Sistem operasi adalah Windows 11 Home 10.0.26200 build "
        "26200. Toolchain menggunakan `rustc 1.97.0 (2d8144b78 2026-07-07)` dan "
        "`cargo 1.97.0` pada profil `--release`.",
    )
    add_para(
        doc,
        "Binary dibangun ulang sebelum pengukuran pada commit `75d17fd`. Pengukuran waktu "
        "menggunakan jam dinding PowerShell dari luar proses. Karena proses baru selalu "
        "menambah overhead spawn, EXP-05 memakai 30 run kontrol `aksara -h` sebagai baseline. "
        "Hasil berlaku pada mesin ini dan tidak digeneralisasi.",
    )

    add_subheading(doc, 5, 2, "Hasil Pengujian EXP-01 sampai EXP-05")
    add_para(
        doc,
        "`cargo test --release` menghasilkan 46/46 test lulus, 0 gagal, dan 0 diabaikan. "
        "Angka tersebut mengonfirmasi perilaku yang diuji, tetapi tidak membuktikan keamanan "
        "sistem secara keseluruhan. Tabel 5.1 merangkum hasil per kelompok eksperimen.",
    )
    add_table_from_file(
        doc,
        "Tabel 5.1 Skenario dan Hasil Pengujian",
        "11_test_scenarios.md",
        landscape=True,
    )
    add_para(
        doc,
        "EXP-01 menghasilkan 10/10 proses unseal sukses dengan invite identik dan 100% "
        "penolakan passphrase salah. Gambar 5.1 memperlihatkan dua unseal dengan passphrase "
        "dummy yang sama serta satu penolakan. Passphrase tampak karena stdin masih "
        "ter-echo; gambar ini sekaligus bukti visual keterbatasan tersebut.",
    )
    add_figure(
        doc,
        screenshots / "04-verifikasi-vault.png",
        "Gambar 5.1 Verifikasi Vault: Determinisme dan Penolakan",
    )
    add_para(
        doc,
        "EXP-02 mempunyai lima test handshake yang lulus. Selisih waktu berpasangan antara "
        "test handshake dan kontrol 0 test adalah -0,15 ms dengan simpangan baku 1,91 ms "
        "pada n=19. Karena nilai tidak terdeteksi di atas noise metode eksternal, hasil "
        "dilaporkan hanya sebagai batas atas kurang dari 0,86 ms pada 95% CI, bukan nilai "
        "latensi titik.",
    )
    add_para(
        doc,
        "EXP-03 mempunyai sembilan test sesi/framing yang lulus. Overhead tag 16 byte "
        "terukur pada instance vault melalui 108 - 16 salt - 12 nonce - 64 plaintext. "
        "Instance Noise transport tidak diukur langsung. EXP-04 mempunyai sepuluh test yang "
        "lulus; invite LAN-only konsisten 86 karakter dan fingerprint 64 heksadesimal pada "
        "lima sampel.",
    )
    add_para(
        doc,
        "EXP-05 menjalankan 30 cold-start unseal. Waktu end-to-end rata-rata 68,47 ms, "
        "median 64,15 ms, simpangan baku 12,47 ms, minimum 54,53 ms, dan maksimum 106,86 ms. "
        "Kontrol tanpa Argon2id mempunyai rata-rata 20,48 ms. Selisih berpasangan memberi "
        "biaya unseal neto rata-rata 47,99 ms, median 45,08 ms, simpangan baku 11,41 ms, "
        "minimum 26,99 ms, dan maksimum 86,30 ms. Lima vault independen seluruhnya berukuran "
        "108 byte.",
    )
    add_para(
        doc,
        "Tabel 5.2 menggabungkan metrik, metode, baseline, hasil, dan status. Kolom PARTIAL "
        "dipertahankan untuk latensi handshake serta overhead transport.",
    )
    add_table_from_file(
        doc,
        "Tabel 5.2 Parameter Evaluasi dan Hasil Terukur",
        "12_evaluation_parameters.md",
        landscape=True,
    )

    add_subheading(doc, 5, 3, "Analisis dan Diskusi")
    add_para(
        doc,
        "Seluruh ekspektasi correctness dan rejection pada test yang tersedia terpenuhi. "
        "Ukuran vault 108 byte, panjang invite 86 karakter, dan tag 16 byte sesuai "
        "spesifikasi internal, sehingga nilai deterministik tersebut kini mempunyai bukti "
        "empiris selain pembacaan konstanta.",
    )
    add_para(
        doc,
        "Komentar `identity::vault::argon2_params` memperkirakan unlock sekitar 100 ms. "
        "Pada hardware uji, biaya neto yang terukur 47,99 ms dengan n=30 dan simpangan baku "
        "11,41 ms, sedangkan end-to-end 68,47 ms. Karena hanya satu mesin diuji, hasil ini "
        "mengoreksi klaim pada lingkungan tersebut dan tidak berarti komentar selalu keliru "
        "pada mesin lain. Argon2id bersifat memory-hard dan performanya bergantung hardware "
        "[[CITE:rfc9106,biryukov2016argon2]].",
    )
    add_para(
        doc,
        "Handshake berada di bawah resolusi metode, sedangkan derivasi kunci vault sekitar "
        "puluhan milidetik. Data mengindikasikan biaya kriptografi pada siklus pemakaian "
        "didominasi unseal, tetapi tidak mendukung pernyataan bahwa handshake tidak memakan "
        "waktu. Pengukuran presisi memerlukan timer internal atau benchmark harness.",
    )
    add_para(
        doc,
        "Pengukuran eksternal menyisakan variasi spawn proses. Overhead instance Noise "
        "transport dan RSS puncak Argon2id tidak observable dari metode ini. Keterbatasan "
        "tersebut menjadi bagian kesimpulan, bukan angka yang diestimasi.",
    )

    # BAB VI
    add_chapter(
        doc,
        6,
        "Bab ini menjawab tiga rumusan masalah, merangkum batas penelitian, dan menyusun "
        "saran yang tertaut langsung pada gap atau temuan.",
    )
    add_subheading(doc, 6, 1, "Kesimpulan")
    add_para(
        doc,
        "Rumusan masalah pertama terjawab penuh pada level implementasi source. AKSARA "
        "menggunakan tujuh komponen inti dengan suite "
        "`Noise_IK_25519_ChaChaPoly_BLAKE2s` yang terkonfirmasi literal. Ed25519 hadir "
        "sebagai keypair identitas dan bahan fingerprint, tetapi tidak dipakai untuk "
        "sign/verify. ChaCha20-Poly1305 dipakai pada tiga konteks tanpa AAD. Sebanyak 46/46 "
        "test correctness dan rejection yang tersedia lulus.",
    )
    add_para(
        doc,
        "Rumusan masalah kedua terjawab penuh secara deskriptif dan pada bagian yang dapat "
        "diuji. Vault 108 byte memakai Argon2id `m=19 MiB`, `t=2`, `p=1`; passphrase salah "
        "ditolak dengan pesan generik; dan 10 unseal menghasilkan invite identik. "
        "Zeroization kuat pada tipe inti tetapi belum merata pada boundary, serta rotasi dan "
        "revokasi tidak ditemukan.",
    )
    add_para(
        doc,
        "Rumusan masalah ketiga terjawab sebagian. Pada laptop Intel Core i7-1165G7, biaya "
        "unseal neto rata-rata 47,99 ms dengan n=30 dan simpangan baku 11,41 ms, sehingga "
        "estimasi sekitar 100 ms pada komentar kode terkoreksi untuk lingkungan tersebut. "
        "Latensi handshake hanya mempunyai batas atas kurang dari 0,86 ms, sedangkan "
        "overhead Noise transport dan RSS belum diukur langsung.",
    )
    add_para(
        doc,
        "Secara sintesis, AKSARA mengimplementasikan rangkaian primitif yang koheren dan "
        "lolos test yang tersedia. Namun, kontak baru masih mengikuti trust-on-first-use "
        "implisit, dan forward secrecy serta identity-hiding tetap `DOCUMENTED_ONLY` karena "
        "tidak diverifikasi oleh test lokal. Kesimpulan ini tidak menyatakan AKSARA aman "
        "secara absolut.",
    )

    add_subheading(doc, 6, 2, "Keterbatasan Penelitian")
    add_para(
        doc,
        "Keterbatasan pertama berasal dari cakupan: dependency eksternal dan CSPRNG OS tidak "
        "diaudit, tidak ada pembuktian formal, side-channel fisik, evaluasi DoS skala luas, "
        "atau skenario multi-device. Penelitian juga tidak mengimplementasikan remediasi dan "
        "menempatkan LAN/Tor sebagai konteks.",
    )
    add_para(
        doc,
        "Keterbatasan kedua berasal dari metode. Timing diukur dari luar proses dan hanya "
        "pada satu unit hardware. Latensi handshake presisi, overhead Noise transport, dan "
        "RSS puncak belum terukur. Test suite tidak mempunyai known-answer test, test vector "
        "standar, maupun rejection test khusus untuk frame transport yang dimodifikasi.",
    )
    add_para(
        doc,
        "Keterbatasan ketiga melekat pada objek: tidak ada rotasi/ratcheting/revokasi, "
        "fallback offline, atau persistensi contact store end-to-end melalui CLI. mDNS "
        "membocorkan presence/fingerprint plaintext, dan passphrase stdin masih ter-echo. "
        "Akibatnya, validitas kesimpulan berada pada perilaku observable dan pemetaan source, "
        "bukan pembuktian keamanan kriptografis.",
    )

    add_subheading(doc, 6, 3, "Saran")
    add_para(
        doc,
        "Saran disusun dari gap G1–G5 dan temuan T1–T7, sehingga tidak ada butir yang "
        "ditambahkan tanpa asal-usul.",
    )
    suggestions = [
        "Melakukan verifikasi formal terhadap instansiasi Noise_IK AKSARA dengan Noise "
        "Explorer atau alat setara untuk menutup G1.",
        "Menambahkan benchmark harness internal, misalnya Criterion, agar latensi handshake "
        "dan overhead ciphertext transport dapat diukur presisi.",
        "Mengevaluasi overhead protokol pada LAN dan Tor nyata untuk menutup G3.",
        "Mengulang benchmark Argon2id pada beberapa kelas hardware untuk menguji apakah "
        "koreksi timing bersifat umum.",
        "Menambahkan mekanisme rotasi atau ratcheting kunci sesi untuk mengurangi dampak "
        "jangka panjang G2/T7.",
        "Menambahkan rejection test terhadap frame atau ciphertext transport yang sengaja "
        "dimodifikasi.",
        "Memperjelas sinyal UI saat sesi berhenti akibat kegagalan dekripsi dibandingkan "
        "penutupan normal untuk menindaklanjuti T5.",
        "Menerapkan hardening permission pada vault dan state Tor serta meninjau "
        "`FS_MISTRUST_DISABLE_PERMISSIONS_CHECKS` untuk T4/T6.",
        "Menyusun handshake kontak yang lebih terstruktur agar verifikasi tidak sepenuhnya "
        "bergantung pada fingerprint manual, menindaklanjuti G4/T1/T3.",
    ]
    for index, item in enumerate(suggestions):
        add_lettered_item(doc, item, index)
    add_para(
        doc,
        "Prioritas awal disarankan pada mekanisme rotasi/ratcheting dan proses kontak, karena "
        "keduanya menyasar temuan berdampak tinggi. Benchmark harness serta rejection test "
        "dapat dilakukan lebih dahulu sebagai peningkatan evaluasi berbiaya implementasi "
        "relatif rendah. Seluruh saran masih berupa usulan dan belum diimplementasikan.",
    )


def add_bibliography(doc: Document):
    missing = CITES.uncited()
    if missing:
        raise AssertionError(f"Referensi belum disitasi: {missing}")
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("DAFTAR PUSTAKA")
    for run in p.runs:
        set_run_font(run, bold=True)
    for key in CITES.order:
        number = CITES.number[key]
        text = ieee_entry(number, BIB[key])
        p = doc.add_paragraph(style="Bibliography")
        # Render markdown italics from formatter.
        add_markup_runs(p, text)


def audit_text(doc: Document):
    full_text = "\n".join(p.text for p in doc.paragraphs)
    forbidden = [
        "CARAKA",
        "CLAMP",
        "Compact Lightweight Authenticated Mesh Protocol",
        "Ascon-MAC",
        "Ascon-Hash256",
        "Epidemic Sync",
        "Controlled Flooding",
    ]
    found = [term for term in forbidden if term.lower() in full_text.lower()]
    if found:
        raise AssertionError(f"Istilah proyek lain ditemukan: {found}")
    banned_title = ["Tanpa Jejak", "Anti-Sadap", "Sepenuhnya Anonim", "Terbukti Aman"]
    if any(term.lower() in TITLE.lower() for term in banned_title):
        raise AssertionError("Judul memuat frasa terlarang")
    if "AKSARA menggunakan tanda tangan digital Ed25519" in full_text:
        raise AssertionError("Klaim Ed25519 terlarang ditemukan")
    if "AKSARA terbukti memiliki forward secrecy" in full_text:
        raise AssertionError("Klaim forward secrecy terlarang ditemukan")
    if "[[CITE:" in full_text:
        raise AssertionError("Placeholder sitasi tersisa")
    if "[PERLU KONFIRMASI]" in full_text:
        raise AssertionError("Marker konfirmasi tak terduga")


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    if not WORKING_TEMPLATE.exists():
        shutil.copy2(REFERENCE, WORKING_TEMPLATE)
    doc = Document(str(WORKING_TEMPLATE))
    clear_document_body(doc)
    configure_styles(doc)
    numbering = create_numbering(doc)
    doc.core_properties.title = TITLE
    doc.core_properties.author = "; ".join(name for name, _ in MEMBERS)
    doc.core_properties.subject = "Tugas Mata Kuliah Implementasi Kriptografi"
    doc.core_properties.comments = "Dihasilkan dari content pack mini-TA AKSARA."
    add_cover(doc)
    add_front_matter(doc)
    add_main_section(doc)
    build_content(doc, numbering)
    add_bibliography(doc)
    continue_main_page_numbering(doc)
    set_update_fields(doc)
    audit_text(doc)
    if len(CITES.order) != 40:
        raise AssertionError(f"Jumlah referensi tersitasi {len(CITES.order)}, bukan 40")
    doc.save(str(FINAL))
    print(f"Saved={FINAL}")
    print(f"Citations={len(CITES.order)}")
    print(f"Paragraphs={len(doc.paragraphs)} Tables={len(doc.tables)} InlineShapes={len(doc.inline_shapes)}")


if __name__ == "__main__":
    main()
